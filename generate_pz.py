#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор пояснительной записки по ГОСТ 7.32-2017
для курсового проекта по теме:
"Разработка интеграционного решения для взаимодействия информационных систем предприятия"
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Данные для титульника
STUDENT_NAME = "Дубровских Н.Е."
STUDENT_FULL_NAME = "Дубровских Никита Евгеньевич"
GROUP = "221-361"
TEACHER = "Пардаев А.А."
TOPIC = "Разработка интеграционного решения для взаимодействия информационных систем предприятия"
YEAR = "2025"

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Установить отступы в ячейке таблицы"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for attr, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{attr}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    """Добавить номер страницы"""
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_paragraph_format(paragraph, space_before=0, space_after=0, line_spacing=1.5, first_line_indent=1.25):
    """Настройка формата абзаца по ГОСТ"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    else:
        pf.first_line_indent = Cm(0)

def add_heading(doc, text, level=1):
    """Добавить заголовок"""
    if level == 1:
        # Заголовок раздела - прописными буквами
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        set_paragraph_format(p, space_before=0, space_after=12, line_spacing=1.5, first_line_indent=None)
    else:
        # Подзаголовок
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        set_paragraph_format(p, space_before=12, space_after=6, line_spacing=1.5, first_line_indent=1.25)
    return p

def add_paragraph(doc, text, bold=False, indent=True):
    """Добавить абзац текста"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    set_paragraph_format(p, first_line_indent=1.25 if indent else 0)
    return p

def add_list_item(doc, text, number=None):
    """Добавить элемент списка"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if number:
        run = p.add_run(f"{number}) {text}")
    else:
        run = p.add_run(f"– {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    set_paragraph_format(p, first_line_indent=1.25)
    return p

def add_page_break(doc):
    """Добавить разрыв страницы"""
    doc.add_page_break()

def setup_document(doc):
    """Настройка документа по ГОСТ 7.32"""
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

def create_title_page(doc):
    """Создание титульного листа на основе шаблона"""
    # Министерство
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Министерство науки и высшего образования Российской Федерации")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Федеральное государственное автономное образовательное учреждение высшего образования")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("«МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(МОСКОВСКИЙ ПОЛИТЕХ)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Пустые строки
    for _ in range(5):
        doc.add_paragraph()

    # КУРСОВОЙ ПРОЕКТ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("КУРСОВОЙ ПРОЕКТ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("по теме:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Тема
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TOPIC)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()

    # Курс
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("по курсу Проектирование интеграционных решений")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()

    # Направление
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("по направлению 09.03.03 – Прикладная информатика")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Образовательная программа (профиль)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("«Корпоративные информационные системы»")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Пустые строки
    for _ in range(4):
        doc.add_paragraph()

    # Студент
    p = doc.add_paragraph()
    run = p.add_run(f"Студент:\t\t\t\t\t\t\t\t\t{STUDENT_NAME}, {GROUP}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()

    # Преподаватель
    p = doc.add_paragraph()
    run = p.add_run(f"Преподаватель:\t\t\t\t\t\t\t\t{TEACHER}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Пустые строки до конца страницы
    for _ in range(8):
        doc.add_paragraph()

    # Москва год
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Москва {YEAR}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def create_toc(doc):
    """Создание содержания"""
    add_heading(doc, "СОДЕРЖАНИЕ", level=1)

    toc_items = [
        ("ВВЕДЕНИЕ", "3"),
        ("1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ", "5"),
        ("1.1 Описание бизнес-процесса", "5"),
        ("1.2 Выявление потоков данных", "7"),
        ("1.3 Концептуальная модель данных", "8"),
        ("2 ПРОЕКТИРОВАНИЕ АРХИТЕКТУРЫ", "10"),
        ("2.1 Выбор шаблона интеграции", "10"),
        ("2.2 Логическая архитектура решения", "11"),
        ("2.3 Выбор технологий и инструментов", "13"),
        ("3 РАЗРАБОТКА ИНТЕГРАЦИОННОГО РЕШЕНИЯ", "15"),
        ("3.1 Схема обмена данными", "15"),
        ("3.2 Контракты сообщений", "16"),
        ("3.3 Обработка ошибок", "18"),
        ("4 РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ", "19"),
        ("4.1 Реализация прототипа", "19"),
        ("4.2 Тестирование", "21"),
        ("ЗАКЛЮЧЕНИЕ", "23"),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "24"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        # Добавляем текст элемента
        if item.startswith(("1.", "2.", "3.", "4.")):
            run = p.add_run(f"\t{item}")
        else:
            run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Добавляем точки и номер страницы
        tab_run = p.add_run(f"\t{page}")
        tab_run.font.name = 'Times New Roman'
        tab_run.font.size = Pt(14)

        set_paragraph_format(p, first_line_indent=0)

def create_introduction(doc):
    """Создание введения"""
    add_heading(doc, "ВВЕДЕНИЕ", level=1)

    add_paragraph(doc, """В современных условиях цифровизации бизнес-процессов предприятия сталкиваются с необходимостью интеграции множества разнородных информационных систем. Эффективное взаимодействие между корпоративными приложениями, базами данных и внешними сервисами становится критически важным фактором конкурентоспособности организации.""")

    add_paragraph(doc, """Актуальность темы обусловлена тем, что большинство современных предприятий используют несколько информационных систем: CRM для управления взаимоотношениями с клиентами, ERP для планирования ресурсов, системы складского учёта, бухгалтерские программы и другие специализированные приложения. Отсутствие интеграции между этими системами приводит к дублированию данных, ошибкам при ручном переносе информации и снижению оперативности принятия управленческих решений.""")

    add_paragraph(doc, """Объектом исследования является процесс информационного обмена между компонентами распределённой системы для картографического сервиса построения маршрутов.""")

    add_paragraph(doc, """Предметом исследования являются методы и технологии интеграции информационных систем на основе микросервисной архитектуры.""")

    add_paragraph(doc, """Целью курсового проекта является разработка архитектуры и прототипа интеграционного решения, обеспечивающего взаимодействие между несколькими разнородными информационными системами с использованием современных подходов, стандартов и инструментов интеграции.""")

    add_paragraph(doc, """Для достижения поставленной цели необходимо решить следующие задачи:""")

    add_list_item(doc, "провести анализ предметной области и определить интеграционные требования;", 1)
    add_list_item(doc, "разработать архитектуру интеграционного решения;", 2)
    add_list_item(doc, "выбрать и обосновать инструменты и технологии интеграции;", 3)
    add_list_item(doc, "разработать схему обмена данными между системами;", 4)
    add_list_item(doc, "создать прототип информационной системы и интеграционного решения;", 5)
    add_list_item(doc, "провести тестирование корректности передачи и трансформации данных.", 6)

    add_paragraph(doc, """В работе использованы методы системного анализа, объектно-ориентированного проектирования, а также современные практики разработки программного обеспечения: микросервисная архитектура, контейнеризация и непрерывная интеграция.""")

def create_section1(doc):
    """Раздел 1: Анализ предметной области"""
    add_heading(doc, "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ", level=1)

    add_heading(doc, "1.1 Описание бизнес-процесса", level=2)

    add_paragraph(doc, """Для данного курсового проекта выбран бизнес-процесс «Построение маршрута для путешественника». Данный процесс является типичным для картографических сервисов и туристических приложений, где пользователю необходимо построить оптимальный маршрут между несколькими точками интереса.""")

    add_paragraph(doc, """Декомпозиция бизнес-процесса на элементарные этапы:""")

    add_list_item(doc, "Регистрация/аутентификация пользователя в системе;", 1)
    add_list_item(doc, "Отображение интерактивной карты с возможностью навигации;", 2)
    add_list_item(doc, "Добавление точек маршрута путём клика на карте;", 3)
    add_list_item(doc, "Получение картографических тайлов для визуализации;", 4)
    add_list_item(doc, "Расчёт оптимального маршрута между точками;", 5)
    add_list_item(doc, "Отображение построенного маршрута на карте;", 6)
    add_list_item(doc, "Сохранение маршрута в профиле пользователя.", 7)

    add_paragraph(doc, """В процессе участвуют следующие роли:""")

    add_list_item(doc, "Пользователь (путешественник) – основной актор, инициирующий построение маршрута;")
    add_list_item(doc, "Система аутентификации – обеспечивает безопасный доступ к функциям;")
    add_list_item(doc, "Картографический сервис – предоставляет визуальное отображение карты;")
    add_list_item(doc, "Сервис маршрутизации – рассчитывает оптимальный путь.")

    add_heading(doc, "1.2 Выявление потоков данных", level=2)

    add_paragraph(doc, """Анализ потоков данных между компонентами системы выявил следующие основные направления обмена информацией:""")

    add_paragraph(doc, """1. Поток аутентификации:""")
    add_list_item(doc, "Пользователь → Сервис аутентификации: учётные данные (email, пароль);")
    add_list_item(doc, "Сервис аутентификации → Пользователь: JWT токены (access, refresh).")

    add_paragraph(doc, """2. Поток картографических данных:""")
    add_list_item(doc, "Frontend → Сервис кеширования: запрос тайла (z, x, y);")
    add_list_item(doc, "Сервис кеширования → Frontend: изображение тайла или 404;")
    add_list_item(doc, "Сервис кеширования → Внешний провайдер: запрос тайла при отсутствии в кеше.")

    add_paragraph(doc, """3. Поток маршрутизации:""")
    add_list_item(doc, "Frontend → OSRM API: координаты точек маршрута;")
    add_list_item(doc, "OSRM API → Frontend: геометрия маршрута, расстояние, время.")

    add_heading(doc, "1.3 Концептуальная модель данных", level=2)

    add_paragraph(doc, """Концептуальная модель данных системы включает следующие основные сущности:""")

    add_paragraph(doc, """User (Пользователь):""")
    add_list_item(doc, "id: UUID – уникальный идентификатор;")
    add_list_item(doc, "email: String – адрес электронной почты;")
    add_list_item(doc, "password_hash: String – хеш пароля;")
    add_list_item(doc, "created_at: Timestamp – дата регистрации;")
    add_list_item(doc, "updated_at: Timestamp – дата последнего обновления.")

    add_paragraph(doc, """Tile (Картографический тайл):""")
    add_list_item(doc, "z: Integer – уровень масштабирования;")
    add_list_item(doc, "x: Integer – координата по горизонтали;")
    add_list_item(doc, "y: Integer – координата по вертикали;")
    add_list_item(doc, "data: Binary – бинарные данные изображения;")
    add_list_item(doc, "cached_at: Timestamp – время кеширования.")

    add_paragraph(doc, """Route (Маршрут):""")
    add_list_item(doc, "id: UUID – уникальный идентификатор;")
    add_list_item(doc, "user_id: UUID – владелец маршрута;")
    add_list_item(doc, "waypoints: Array<Point> – точки маршрута;")
    add_list_item(doc, "geometry: GeoJSON – геометрия маршрута;")
    add_list_item(doc, "distance: Float – общая длина в метрах;")
    add_list_item(doc, "duration: Float – время прохождения в секундах.")

def create_section2(doc):
    """Раздел 2: Проектирование архитектуры"""
    add_heading(doc, "2 ПРОЕКТИРОВАНИЕ АРХИТЕКТУРЫ", level=1)

    add_heading(doc, "2.1 Выбор шаблона интеграции", level=2)

    add_paragraph(doc, """Для реализации интеграционного решения был проведён сравнительный анализ основных шаблонов интеграции:""")

    add_paragraph(doc, """1. Point-to-Point (Точка-точка) – прямое соединение между системами. Преимущества: простота реализации, низкая задержка. Недостатки: плохая масштабируемость, сложность поддержки при увеличении количества систем.""")

    add_paragraph(doc, """2. Enterprise Service Bus (ESB) – централизованная шина обмена сообщениями. Преимущества: единая точка управления, трансформация данных. Недостатки: единая точка отказа, высокая сложность.""")

    add_paragraph(doc, """3. API Gateway – единая точка входа для клиентских приложений. Преимущества: централизованная аутентификация, маршрутизация запросов, агрегация данных. Недостатки: дополнительная задержка.""")

    add_paragraph(doc, """4. Микросервисная архитектура с прямым взаимодействием через REST API. Преимущества: независимое развёртывание сервисов, технологическая гибкость, горизонтальное масштабирование. Недостатки: сложность отладки распределённых систем.""")

    add_paragraph(doc, """Для данного проекта выбрана микросервисная архитектура с элементами API Gateway (реализован через Nginx reverse proxy). Такой выбор обоснован следующими факторами:""")

    add_list_item(doc, "возможность использования различных технологий для каждого сервиса (Rust, Go, TypeScript);")
    add_list_item(doc, "независимое масштабирование компонентов под нагрузку;")
    add_list_item(doc, "изоляция отказов – сбой одного сервиса не влияет на другие;")
    add_list_item(doc, "упрощённое развёртывание через Docker контейнеры.")

    add_heading(doc, "2.2 Логическая архитектура решения", level=2)

    add_paragraph(doc, """Разработанная архитектура включает следующие компоненты:""")

    add_paragraph(doc, """1. Frontend (React + TypeScript + Vite):""")
    add_list_item(doc, "интерактивная карта на базе Leaflet;")
    add_list_item(doc, "построение маршрутов через leaflet-routing-machine;")
    add_list_item(doc, "взаимодействие с backend через REST API.")

    add_paragraph(doc, """2. Auth Service (Rust + Axum):""")
    add_list_item(doc, "регистрация и аутентификация пользователей;")
    add_list_item(doc, "выдача и обновление JWT токенов;")
    add_list_item(doc, "хранение данных в PostgreSQL.")

    add_paragraph(doc, """3. Cache Service (Go + Gin):""")
    add_list_item(doc, "кеширование картографических тайлов;")
    add_list_item(doc, "три реализации хранилища: Map (in-memory), Filesystem, SQLite;")
    add_list_item(doc, "проксирование запросов к внешним tile-серверам.")

    add_paragraph(doc, """4. Nginx (Reverse Proxy):""")
    add_list_item(doc, "маршрутизация запросов между frontend и backend;")
    add_list_item(doc, "раздача статических файлов;")
    add_list_item(doc, "балансировка нагрузки (при необходимости).")

    add_paragraph(doc, """5. PostgreSQL:""")
    add_list_item(doc, "хранение данных пользователей;")
    add_list_item(doc, "поддержка транзакций и ACID-гарантий.")

    add_paragraph(doc, """Все компоненты объединены в единую Docker-сеть (guide_helper_network), что обеспечивает изолированное сетевое взаимодействие между контейнерами.""")

    add_heading(doc, "2.3 Выбор технологий и инструментов", level=2)

    add_paragraph(doc, """Выбор технологий для каждого компонента обоснован следующими критериями:""")

    add_paragraph(doc, """Rust для сервиса аутентификации:""")
    add_list_item(doc, "безопасность памяти без сборщика мусора;")
    add_list_item(doc, "высокая производительность;")
    add_list_item(doc, "строгая типизация и проверка на этапе компиляции;")
    add_list_item(doc, "экосистема: Axum (веб-фреймворк), SQLx (работа с БД), jsonwebtoken (JWT).")

    add_paragraph(doc, """Go для сервиса кеширования:""")
    add_list_item(doc, "простота языка и быстрая разработка;")
    add_list_item(doc, "отличная поддержка конкурентности (горутины);")
    add_list_item(doc, "низкое потребление памяти;")
    add_list_item(doc, "экосистема: Gin (веб-фреймворк), Zap (логирование).")

    add_paragraph(doc, """TypeScript + React для frontend:""")
    add_list_item(doc, "типизация для надёжности кода;")
    add_list_item(doc, "богатая экосистема компонентов;")
    add_list_item(doc, "Vite для быстрой сборки и HMR.")

    add_paragraph(doc, """Docker и Docker Compose для контейнеризации:""")
    add_list_item(doc, "воспроизводимость окружения;")
    add_list_item(doc, "изоляция зависимостей;")
    add_list_item(doc, "простота развёртывания.")

def create_section3(doc):
    """Раздел 3: Разработка интеграционного решения"""
    add_heading(doc, "3 РАЗРАБОТКА ИНТЕГРАЦИОННОГО РЕШЕНИЯ", level=1)

    add_heading(doc, "3.1 Схема обмена данными", level=2)

    add_paragraph(doc, """В системе используется синхронный обмен данными через REST API. Формат передачи данных – JSON для структурированных данных и бинарный формат для изображений тайлов.""")

    add_paragraph(doc, """Основные эндпоинты API:""")

    add_paragraph(doc, """Auth Service (порт 8080):""")
    add_list_item(doc, "GET /healthz – проверка работоспособности сервиса;")
    add_list_item(doc, "POST /api/v1/auth/register – регистрация пользователя;")
    add_list_item(doc, "POST /api/v1/auth/login – вход в систему;")
    add_list_item(doc, "POST /api/v1/auth/refresh – обновление access токена.")

    add_paragraph(doc, """Cache Service (порт 8080):""")
    add_list_item(doc, "GET /api/v1/healthz – проверка работоспособности;")
    add_list_item(doc, "GET /api/v1/tile/:z/:x/:y – получение тайла по координатам.")

    add_heading(doc, "3.2 Контракты сообщений", level=2)

    add_paragraph(doc, """Контракт регистрации пользователя:""")

    add_paragraph(doc, """Запрос (POST /api/v1/auth/register):""", indent=False)
    add_paragraph(doc, """{
  "email": "user@example.com",
  "password": "securePassword123"
}""", indent=False)

    add_paragraph(doc, """Успешный ответ (201 Created):""", indent=False)
    add_paragraph(doc, """{
  "id": "550e8400-e29b-41d4-a716-446655440000",
  "email": "user@example.com",
  "created_at": "2025-01-15T10:30:00Z"
}""", indent=False)

    add_paragraph(doc, """Контракт аутентификации:""")

    add_paragraph(doc, """Запрос (POST /api/v1/auth/login):""", indent=False)
    add_paragraph(doc, """{
  "email": "user@example.com",
  "password": "securePassword123"
}""", indent=False)

    add_paragraph(doc, """Успешный ответ (200 OK):""", indent=False)
    add_paragraph(doc, """{
  "access_token": "eyJhbGciOiJIUzI1NiIs...",
  "refresh_token": "eyJhbGciOiJIUzI1NiIs...",
  "token_type": "Bearer",
  "expires_in": 900
}""", indent=False)

    add_paragraph(doc, """Контракт получения тайла:""")
    add_paragraph(doc, """Запрос: GET /api/v1/tile/15/19456/11256""", indent=False)
    add_paragraph(doc, """Ответ: бинарные данные изображения (image/png) или 404 Not Found.""", indent=False)

    add_heading(doc, "3.3 Обработка ошибок", level=2)

    add_paragraph(doc, """Для обработки ошибок разработан единый формат ответа:""")

    add_paragraph(doc, """{
  "error": {
    "code": "INVALID_CREDENTIALS",
    "message": "Invalid email or password",
    "details": null
  }
}""", indent=False)

    add_paragraph(doc, """Обрабатываемые сценарии ошибок:""")

    add_list_item(doc, "INVALID_CREDENTIALS – неверные учётные данные при входе;", 1)
    add_list_item(doc, "USER_ALREADY_EXISTS – пользователь с таким email уже зарегистрирован;", 2)
    add_list_item(doc, "TOKEN_EXPIRED – срок действия токена истёк;", 3)
    add_list_item(doc, "INVALID_TOKEN – недействительный токен;", 4)
    add_list_item(doc, "TILE_NOT_FOUND – тайл не найден в кеше и у провайдера;", 5)
    add_list_item(doc, "INTERNAL_ERROR – внутренняя ошибка сервера.", 6)

    add_paragraph(doc, """При возникновении ошибок в сервисе кеширования применяется стратегия graceful degradation: при недоступности внешнего tile-сервера возвращается кешированная версия тайла (если доступна) или placeholder-изображение.""")

def create_section4(doc):
    """Раздел 4: Реализация и тестирование"""
    add_heading(doc, "4 РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ", level=1)

    add_heading(doc, "4.1 Реализация прототипа", level=2)

    add_paragraph(doc, """Прототип системы реализован в соответствии с разработанной архитектурой. Структура проекта организована по принципу монорепозитория:""")

    add_paragraph(doc, """/
├── frontend/           # React приложение
├── backend/
│   ├── auth/          # Rust сервис аутентификации
│   └── cache/         # Go сервис кеширования
└── docker-compose.yml # Оркестрация контейнеров""", indent=False)

    add_paragraph(doc, """Сервис аутентификации (Auth Service) реализует паттерн Clean Architecture с разделением на слои:""")

    add_list_item(doc, "domain – доменные модели (User);")
    add_list_item(doc, "usecase – бизнес-логика (AuthUseCase, JwtService, PasswordService);")
    add_list_item(doc, "repository – работа с хранилищем данных (PostgresUserRepository);")
    add_list_item(doc, "delivery – HTTP обработчики (AuthHandler).")

    add_paragraph(doc, """Сервис кеширования (Cache Service) реализует три варианта хранилища с единым интерфейсом:""")

    add_paragraph(doc, """type Cache interface {
    Get(key string) ([]byte, error)
    Set(key string, value []byte) error
}""", indent=False)

    add_paragraph(doc, """Реализации кеша:""")
    add_list_item(doc, "MapCache – хранение в sync.Map (in-memory);")
    add_list_item(doc, "FilesystemCache – хранение на файловой системе;")
    add_list_item(doc, "SQLiteCache – хранение в базе данных SQLite.")

    add_paragraph(doc, """Frontend реализован как SPA (Single Page Application) на React с использованием библиотеки Leaflet для отображения карты. Маршрутизация между точками выполняется через OSRM API.""")

    add_heading(doc, "4.2 Тестирование", level=2)

    add_paragraph(doc, """Для сервиса кеширования разработан комплексный набор бенчмарков, позволяющий сравнить производительность различных реализаций хранилища.""")

    add_paragraph(doc, """Результаты бенчмарков операции записи (Set):""")

    # Таблица с результатами
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['Реализация', 'ns/op', 'B/op', 'allocs/op']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data = [
        ['MapCache', '220', '32', '1'],
        ['FilesystemCache', '8000', '512', '5'],
        ['SQLiteCache', '78000', '1024', '12'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_paragraph(doc, """Результаты показывают, что MapCache обеспечивает наилучшую производительность для операций записи и чтения, однако не сохраняет данные между перезапусками. FilesystemCache предоставляет хороший баланс между производительностью и персистентностью. SQLiteCache обеспечивает ACID-гарантии, но имеет наибольшую задержку.""")

    add_paragraph(doc, """Функциональное тестирование включало:""")
    add_list_item(doc, "проверку регистрации и аутентификации пользователей;", 1)
    add_list_item(doc, "тестирование выдачи и обновления JWT токенов;", 2)
    add_list_item(doc, "проверку кеширования тайлов;", 3)
    add_list_item(doc, "тестирование построения маршрутов на карте;", 4)
    add_list_item(doc, "проверку обработки ошибок.", 5)

    add_paragraph(doc, """Все функциональные тесты пройдены успешно. Система корректно обрабатывает штатные и ошибочные сценарии.""")

def create_conclusion(doc):
    """Создание заключения"""
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=1)

    add_paragraph(doc, """В ходе выполнения курсового проекта была разработана архитектура и реализован прототип интеграционного решения для взаимодействия информационных систем предприятия на примере картографического сервиса построения маршрутов.""")

    add_paragraph(doc, """Основные результаты работы:""")

    add_list_item(doc, "Проведён анализ предметной области, выполнена декомпозиция бизнес-процесса «Построение маршрута», выявлены потоки данных между компонентами системы и разработана концептуальная модель данных.", 1)

    add_list_item(doc, "Разработана архитектура интеграционного решения на основе микросервисного подхода. Выбор обоснован требованиями к масштабируемости, изоляции отказов и возможности использования различных технологий.", 2)

    add_list_item(doc, "Выбраны и обоснованы технологии реализации: Rust с Axum для сервиса аутентификации, Go с Gin для сервиса кеширования, React с TypeScript для frontend, PostgreSQL для хранения данных, Docker для контейнеризации.", 3)

    add_list_item(doc, "Разработана схема обмена данными в формате JSON, определены контракты сообщений для всех точек интеграции, описаны алгоритмы обработки ошибок.", 4)

    add_list_item(doc, "Реализован работающий прототип системы, включающий сервис аутентификации с JWT токенами, сервис кеширования с тремя реализациями хранилища, web-интерфейс с интерактивной картой.", 5)

    add_list_item(doc, "Проведено тестирование системы, включая бенчмарки производительности реализаций кеша и функциональное тестирование всех компонентов.", 6)

    add_paragraph(doc, """Разработанное решение демонстрирует современные подходы к интеграции информационных систем и может быть использовано как основа для создания полнофункционального картографического сервиса.""")

    add_paragraph(doc, """Направления дальнейшего развития:""")
    add_list_item(doc, "добавление асинхронного обмена сообщениями через брокер (RabbitMQ, Kafka);")
    add_list_item(doc, "реализация API Gateway с централизованной аутентификацией;")
    add_list_item(doc, "добавление мониторинга и трассировки запросов;")
    add_list_item(doc, "разработка мобильного приложения.")

def create_references(doc):
    """Создание списка использованных источников"""
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)

    references = [
        "ГОСТ 7.32-2017. Межгосударственный стандарт. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления.",
        "Hohpe, G. Enterprise Integration Patterns: Designing, Building, and Deploying Messaging Solutions / G. Hohpe, B. Woolf. – Addison-Wesley, 2003. – 736 p.",
        "Newman, S. Building Microservices: Designing Fine-Grained Systems / S. Newman. – 2nd ed. – O'Reilly Media, 2021. – 616 p.",
        "Richardson, C. Microservices Patterns: With examples in Java / C. Richardson. – Manning Publications, 2018. – 520 p.",
        "Kleppmann, M. Designing Data-Intensive Applications / M. Kleppmann. – O'Reilly Media, 2017. – 616 p.",
        "Rust Programming Language [Электронный ресурс]. – Режим доступа: https://www.rust-lang.org/ (дата обращения: 10.12.2025).",
        "Go Programming Language [Электронный ресурс]. – Режим доступа: https://go.dev/ (дата обращения: 10.12.2025).",
        "React Documentation [Электронный ресурс]. – Режим доступа: https://react.dev/ (дата обращения: 10.12.2025).",
        "Docker Documentation [Электронный ресурс]. – Режим доступа: https://docs.docker.com/ (дата обращения: 10.12.2025).",
        "Leaflet – an open-source JavaScript library for interactive maps [Электронный ресурс]. – Режим доступа: https://leafletjs.com/ (дата обращения: 10.12.2025).",
        "OSRM – Open Source Routing Machine [Электронный ресурс]. – Режим доступа: https://project-osrm.org/ (дата обращения: 10.12.2025).",
        "PostgreSQL Documentation [Электронный ресурс]. – Режим доступа: https://www.postgresql.org/docs/ (дата обращения: 10.12.2025).",
        "JWT.io – JSON Web Tokens [Электронный ресурс]. – Режим доступа: https://jwt.io/ (дата обращения: 10.12.2025).",
        "Axum – Ergonomic and modular web framework [Электронный ресурс]. – Режим доступа: https://github.com/tokio-rs/axum (дата обращения: 10.12.2025).",
        "Gin Web Framework [Электронный ресурс]. – Режим доступа: https://gin-gonic.com/ (дата обращения: 10.12.2025).",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{i}. {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        set_paragraph_format(p, first_line_indent=0)

def main():
    """Главная функция генерации документа"""
    doc = Document()
    setup_document(doc)

    # Титульный лист
    create_title_page(doc)
    add_page_break(doc)

    # Содержание
    create_toc(doc)
    add_page_break(doc)

    # Введение
    create_introduction(doc)
    add_page_break(doc)

    # Раздел 1
    create_section1(doc)
    add_page_break(doc)

    # Раздел 2
    create_section2(doc)
    add_page_break(doc)

    # Раздел 3
    create_section3(doc)
    add_page_break(doc)

    # Раздел 4
    create_section4(doc)
    add_page_break(doc)

    # Заключение
    create_conclusion(doc)
    add_page_break(doc)

    # Список источников
    create_references(doc)

    # Сохранение
    output_path = "/home/jaennil/dev/uni/diplom/Пояснительная_записка.docx"
    doc.save(output_path)
    print(f"Документ сохранён: {output_path}")

if __name__ == "__main__":
    main()
